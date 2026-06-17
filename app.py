"""
PatentMind — Patent Portfolio Intelligence Dashboard
============================================================
Run:  streamlit run app.py

Tabs:
  1. Portfolio Overview  — KPIs, filing trends, IPC distribution, status
  2. Technology Landscape — UMAP clustering of all patents
  3. Whitespace Analysis  — IPC coverage gaps + competitor comparison
  4. Patent Q&A Chat      — RAG-powered natural language search
  5. Report Generator     — One-click Claude-written portfolio reports
"""


import os, json, re
from pathlib import Path
from datetime import datetime
from dotenv import load_dotenv
load_dotenv(override=True) # loads .env file automatically
import warnings
warnings.filterwarnings("ignore", category=UserWarning)

import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import numpy as np
import chromadb
from sentence_transformers import SentenceTransformer
import anthropic
import voyageai




# ── PAGE CONFIG ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="PatentIntel — Portfolio Intelligence",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── CUSTOM CSS ────────────────────────────────────────────────────────────────
st.markdown("""
<style>
.metric-card{background:#f8f9fa;border-radius:8px;padding:16px;border-left:3px solid #185FA5}
.badge-granted{background:#E1F5EE;color:#085041;padding:2px 8px;border-radius:12px;font-size:11px}
.badge-pending{background:#FAEEDA;color:#633806;padding:2px 8px;border-radius:12px;font-size:11px}
.badge-expired{background:#FAECE7;color:#712B13;padding:2px 8px;border-radius:12px;font-size:11px}
.section-header{font-size:14px;font-weight:600;color:#1a1a2e;margin:16px 0 8px}
</style>
""", unsafe_allow_html=True)

# ── CONSTANTS ─────────────────────────────────────────────────────────────────
CHROMA_PATH  = "./chroma_db"
COLLECTION   = "patents"
##EMBED_MODEL  = "Snowflake/snowflake-arctic-embed-l-v2.0" ##removing as started using voyage-3
DATA_FILE    = "patents_data.json"

# ── LOADERS (cached) ──────────────────────────────────────────────────────────

@st.cache_resource
def load_chroma():
    client = chromadb.PersistentClient(path=CHROMA_PATH)
    return client.get_collection(COLLECTION)

@st.cache_data
def load_dataframe() -> pd.DataFrame:
    if Path(DATA_FILE).exists():
        with open(DATA_FILE) as f:
            data = json.load(f)
        df = pd.DataFrame(data)
    else:
        # Fallback: pull from ChromaDB metadata
        col = load_chroma()
        res = col.get(include=["metadatas"])
        df  = pd.DataFrame(res["metadatas"])

    # Parse dates
    df["filing_date"] = pd.to_datetime(df["filing_date"], errors="coerce")
    df["year"] = df["filing_date"].dt.year
    df["status"] = df["status"].fillna("unknown").str.lower()
    df["ipc_label"] = df["ipc_label"].fillna("Unknown")
    return df

def get_client():
    key = (
        os.environ.get("ANTHROPIC_API_KEY", "")
        or st.secrets.get("ANTHROPIC_API_KEY", "")
    )
    if not key:
        st.error("⚠️ ANTHROPIC_API_KEY not found. Check your .env file.")
        return None
    return anthropic.Anthropic(api_key=key)

# ── SEMANTIC SEARCH ───────────────────────────────────────────────────────────
def search_patents(query: str, n: int = 8) -> list[dict]:
    import voyageai
    vo = voyageai.Client(api_key=os.environ.get("VOYAGE_API_KEY"))
    col = load_chroma()
    result = vo.embed([query], model="voyage-3", input_type="query")
    q_emb = result.embeddings[0]
    res = col.query(query_embeddings=[q_emb], n_results=n, include=["metadatas","distances"])
    results = []
    for meta, dist in zip(res["metadatas"][0], res["distances"][0]):
        results.append({**meta, "similarity": round(1-dist, 3)})
    return results

# ── CLAUDE CALLS ──────────────────────────────────────────────────────────────
def claude_answer(question: str, patents: list[dict]) -> str:
    client = get_client()
    if not client:
        return "⚠️ Please enter your Anthropic API key in the sidebar."

    context = "\n\n".join([
        f"Patent {i+1}: [{p.get('patent_id','')}] {p.get('title','')}\n"
        f"IPC: {p.get('ipc_label','')} | Filed: {p.get('filing_date','')} | "
        f"Status: {p.get('status','')}\nAbstract: {p.get('abstract','')[:400]}"
        for i, p in enumerate(patents)
    ])

    resp = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=1024,
        system="""You are a senior IP analyst. Answer using ONLY the provided patent context.
Always cite patent IDs in [brackets]. Be concise and strategic.
If information is insufficient, say so clearly.""",
        messages=[{
            "role": "user",
            "content": f"Question: {question}\n\nPatent context:\n{context}"
        }]
    )
    return resp.content[0].text


def claude_report(report_type: str, df: pd.DataFrame) -> str:
    client = get_client()
    if not client:
        return "⚠️ Please enter your Anthropic API key in the sidebar."

    # Build portfolio summary stats
    stats = {
        "total_patents":    len(df),
        "granted":          len(df[df.status=="granted"]),
        "pending":          len(df[df.status=="pending"]),
        "expired":          len(df[df.status=="expired"]),
        "top_ipc":          df["ipc_label"].value_counts().head(5).to_dict(),
        "filing_range":     f"{df.year.min():.0f}–{df.year.max():.0f}" if df.year.notna().any() else "N/A",
        "recent_5yr":       len(df[df.year >= datetime.now().year - 5]) if df.year.notna().any() else 0,
    }

    prompts = {
        "Portfolio Summary": f"""Write a 2-page executive IP portfolio summary based on these metrics:
{json.dumps(stats, indent=2)}
Include: portfolio strength assessment, technology focus areas, filing velocity trend,
strategic recommendations, and areas for investment. Use professional IP analyst language.""",

        "Landscape Report": f"""Write a technology landscape analysis for this patent portfolio:
{json.dumps(stats, indent=2)}
Include: technology distribution analysis, dominant IPC classes and their strategic importance,
emerging technology clusters, portfolio depth vs breadth assessment, and competitive positioning insights.""",

        "Whitespace Analysis": f"""Write a patent whitespace analysis report based on these portfolio metrics:
{json.dumps(stats, indent=2)}
Identify: underrepresented technology areas in this portfolio, strategic filing opportunities,
technology domains where competitors likely hold advantage, and recommendations for portfolio expansion.
Structure as: Executive Summary → Gap Analysis → Opportunity Map → Recommended Actions.""",

        "M&A Diligence Summary": f"""Write an IP due diligence assessment template based on this portfolio:
{json.dumps(stats, indent=2)}
Cover: portfolio size and quality metrics, geographic coverage assessment (infer from data),
technology coverage breadth, potential freedom-to-operate concerns, key value drivers,
risk factors, and integration recommendations for M&A context.""",

        "Maintenance Recommendations": f"""Write patent maintenance strategy recommendations for this portfolio:
{json.dumps(stats, indent=2)}
Address: which status categories need immediate attention, cost optimization strategies,
how to evaluate which patents to maintain vs abandon, and a framework for ongoing portfolio pruning.""",
    }

    resp = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=2000,
        system="You are a senior IP counsel and patent analytics expert writing professional IP strategy reports.",
        messages=[{"role": "user", "content": prompts.get(report_type, prompts["Portfolio Summary"])}]
    )
    return resp.content[0].text


def claude_cluster_labels(cluster_texts: dict) -> dict:
    """Ask Claude to label technology clusters from their representative patents."""
    client = get_client()
    if not client:
        return {k: f"Cluster {k}" for k in cluster_texts}

    prompt = "Label each patent cluster with a 2-4 word technology theme:\n\n"
    for cid, titles in cluster_texts.items():
        prompt += f"Cluster {cid}: {', '.join(titles[:3])}\n"
    prompt += "\nRespond ONLY as JSON: {\"0\": \"label\", \"1\": \"label\", ...}"

    resp = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=300,
        messages=[{"role": "user", "content": prompt}]
    )
    try:
        return json.loads(resp.content[0].text)
    except Exception:
        return {k: f"Cluster {k}" for k in cluster_texts}

# ── SIDEBAR ───────────────────────────────────────────────────────────────────
with st.sidebar:
    st.image("https://www.anthropic.com/images/icons/apple-touch-icon.png", width=40)
    st.title("PatentIntel")
    st.caption("Patent Portfolio Intelligence · Powered by Claude")
    st.divider()


    st.divider()
    st.caption("**Portfolio:**")
    try:
        df_sidebar = load_dataframe()
        col = load_chroma()
        st.metric("Total Patents", len(df_sidebar))
        st.metric("Indexed Vectors", col.count())
        granted_pct = len(df_sidebar[df_sidebar.status=="granted"])/len(df_sidebar)*100
        st.metric("Granted", f"{granted_pct:.0f}%")
    except Exception as e:
        st.error(f"⚠️ Failed to load data: {str(e)}")
        st.stop()

    st.divider()
    st.caption("📂 **Data source**")
    data_source = st.selectbox("Input format", ["Pre-loaded (patents_data.json)", "Upload CSV", "Demo mode"])

# ── MAIN CONTENT ──────────────────────────────────────────────────────────────
st.title(" PatentIntel — IP Portfolio Intelligence")
st.caption(f"Portfolio of **{len(load_dataframe())}** patents · Last indexed: {Path('patents_data.json').stat().st_mtime if Path('patents_data.json').exists() else 'N/A'}")

tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "📊 Portfolio Overview",
    "🗺️ Technology Landscape",
    "⬜ Whitespace Analysis",
    "💬 Patent Q&A",
    "📄 Report Generator"
])

# ════════════════════════════════════════════════════════════════════════════════
# TAB 1 — PORTFOLIO OVERVIEW
# ════════════════════════════════════════════════════════════════════════════════
with tab1:
    df = load_dataframe()
    st.markdown('<div class="section-header">Portfolio KPIs</div>', unsafe_allow_html=True)

    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("Total Patents",  len(df))
    c2.metric("Granted",        len(df[df.status=="granted"]),  delta="active assets")
    c3.metric("Pending",        len(df[df.status=="pending"]),  delta="in prosecution")
    c4.metric("Expired",        len(df[df.status=="expired"]),  delta_color="inverse")
    c5.metric("Tech Domains",   df["ipc_label"].nunique())

    st.divider()
    col_a, col_b = st.columns(2)

    with col_a:
        st.markdown('<div class="section-header">Filing Trend by Year</div>', unsafe_allow_html=True)
        yearly = df.groupby("year").size().reset_index(name="count")
        yearly = yearly[yearly.year.between(2000, 2025)]
        fig1 = px.line(yearly, x="year", y="count",
                       markers=True, color_discrete_sequence=["#185FA5"],
                       labels={"year": "Filing Year", "count": "Patents Filed"})
        fig1.update_layout(margin=dict(t=20, b=20), height=280)
        st.plotly_chart(fig1, use_container_width=True)

    with col_b:
        st.markdown('<div class="section-header">Status Distribution</div>', unsafe_allow_html=True)
        status_counts = df["status"].value_counts().reset_index()
        status_counts.columns = ["status", "count"]
        fig2 = px.pie(status_counts, names="status", values="count",
                      color_discrete_map={
                          "granted": "#1D9E75", "pending": "#BA7517",
                          "expired": "#E24B4A", "unknown": "#aaaaaa"
                      }, hole=0.45)
        fig2.update_layout(margin=dict(t=20, b=20), height=280, showlegend=True)
        st.plotly_chart(fig2, use_container_width=True)

    col_c, col_d = st.columns(2)

    with col_c:
        st.markdown('<div class="section-header">Top Technology Domains (IPC)</div>', unsafe_allow_html=True)
        ipc_counts = df["ipc_label"].value_counts().head(12).reset_index()
        ipc_counts.columns = ["domain", "count"]
        fig3 = px.bar(ipc_counts, x="count", y="domain", orientation="h",
                      color="count", color_continuous_scale="Blues",
                      labels={"count": "Patents", "domain": ""})
        fig3.update_layout(margin=dict(t=20, b=20), height=320,
                           coloraxis_showscale=False, yaxis=dict(autorange="reversed"))
        st.plotly_chart(fig3, use_container_width=True)

    with col_d:
        st.markdown('<div class="section-header">Filing Activity Heatmap (Year × Domain)</div>', unsafe_allow_html=True)
        top_ipc = df["ipc_label"].value_counts().head(8).index.tolist()
        heat_df = df[df["ipc_label"].isin(top_ipc) & df["year"].notna()]
        heat_df = heat_df.groupby(["year","ipc_label"]).size().reset_index(name="count")
        if len(heat_df) > 0:
            pivot = heat_df.pivot(index="ipc_label", columns="year", values="count").fillna(0)
            fig4  = px.imshow(pivot, color_continuous_scale="Blues",
                              labels=dict(x="Year", y="", color="Patents"),
                              aspect="auto")
            fig4.update_layout(margin=dict(t=20, b=20), height=320)
            st.plotly_chart(fig4, use_container_width=True)

    st.divider()
    st.markdown('<div class="section-header">Patent List</div>', unsafe_allow_html=True)
    search_q = st.text_input("🔍 Filter table", placeholder="Search title, IPC, status...")
    display_df = df.copy()
    if search_q:
        mask = (
            display_df["title"].str.contains(search_q, case=False, na=False) |
            display_df["ipc_label"].str.contains(search_q, case=False, na=False) |
            display_df["patent_id"].str.contains(search_q, case=False, na=False)
        )
        display_df = display_df[mask]

    st.dataframe(
        display_df[["patent_id","title","ipc_label","filing_date","status","assignee"]].head(200),
        use_container_width=True, height=300,
        column_config={
            "patent_id":   st.column_config.TextColumn("Patent ID", width=120),
            "title":       st.column_config.TextColumn("Title", width=300),
            "ipc_label":   st.column_config.TextColumn("Technology", width=180),
            "filing_date": st.column_config.DateColumn("Filed"),
            "status":      st.column_config.TextColumn("Status", width=90),
        }
    )

# ════════════════════════════════════════════════════════════════════════════════
# TAB 2 — TECHNOLOGY LANDSCAPE (UMAP)
# ════════════════════════════════════════════════════════════════════════════════
with tab2:
    st.markdown("### Technology Landscape Map")
    st.caption("Each dot = one patent. Proximity = similar technology. Clusters reveal natural groupings in your portfolio.")

    if st.button("🗺️ Generate Landscape (may take ~30s first time)", type="primary"):
        with st.spinner("Computing UMAP projections..."):
            try:
                from umap import UMAP
                from sklearn.preprocessing import LabelEncoder

                col    = load_chroma()
                result = col.get(include=["embeddings","metadatas"])
                embs   = np.array(result["embeddings"])
                metas  = result["metadatas"]

                umap_model = UMAP(n_components=2, random_state=42,
                                  n_neighbors=min(15, len(embs)-1), min_dist=0.1)
                coords = umap_model.fit_transform(embs)

                umap_df = pd.DataFrame({
                    "x":          coords[:, 0],
                    "y":          coords[:, 1],
                    "title":      [m.get("title", "")[:60] for m in metas],
                    "patent_id":  [m.get("patent_id", "") for m in metas],
                    "ipc_label":  [m.get("ipc_label", "Unknown") for m in metas],
                    "status":     [m.get("status", "") for m in metas],
                    "filed":      [m.get("filing_date", "") for m in metas],
                })

                fig = px.scatter(
                    umap_df, x="x", y="y",
                    color="ipc_label",
                    hover_data={"patent_id": True, "title": True,
                                "filed": True, "x": False, "y": False},
                    labels={"ipc_label": "Technology Domain"},
                    title="Patent Portfolio — Technology Landscape (UMAP 2D projection)"
                )
                fig.update_traces(marker=dict(size=7, opacity=0.75))
                fig.update_layout(height=550, margin=dict(t=50),
                                  legend=dict(orientation="v", x=1.01))
                st.plotly_chart(fig, use_container_width=True)

                # Claude cluster labels
                if get_client():
                    from sklearn.cluster import KMeans
                    n_clusters = min(8, len(embs))
                    km = KMeans(n_clusters=n_clusters, random_state=42, n_init=10)
                    labels = km.fit_predict(embs)
                    umap_df["cluster"] = labels

                    cluster_titles = {}
                    for cid in range(n_clusters):
                        idx = umap_df[umap_df["cluster"]==cid].index[:3].tolist()
                        cluster_titles[str(cid)] = [umap_df.loc[i,"title"] for i in idx]

                    with st.spinner("Claude labelling clusters..."):
                        cluster_names = claude_cluster_labels(cluster_titles)

                    st.markdown("**AI-identified technology clusters:**")
                    cols = st.columns(4)
                    for i, (cid, name) in enumerate(cluster_names.items()):
                        cnt = (umap_df["cluster"]==int(cid)).sum()
                        cols[i%4].metric(f"Cluster {int(cid)+1}", name, f"{cnt} patents")

            except ImportError:
                st.error("Install umap-learn: `pip install umap-learn`")
            except Exception as e:
                st.error(f"UMAP error: {e}")

# ════════════════════════════════════════════════════════════════════════════════
# TAB 3 — WHITESPACE ANALYSIS
# ════════════════════════════════════════════════════════════════════════════════
with tab3:
    st.markdown("### Whitespace Analysis")
    st.caption("Identify technology areas where your portfolio has gaps — strategic filing opportunities.")

    df = load_dataframe()

    # All major IPC domains in semiconductor + media space
    ALL_DOMAINS = [
        "Semiconductor Devices", "FinFET Transistor", "5G Beamforming",
        "Video Coding", "Audio/Speech Processing", "Image Processing",
        "Neural Architecture", "DRAM Cell", "Cache Coherence",
        "Packet Routing", "Wireless Communication", "LDPC Coding",
        "Power Conversion", "Laser / Photonics", "Data Storage",
        "Computer Vision", "Natural Language Processing", "Edge AI",
        "Autonomous Systems", "Cryptography / Security",
    ]

    your_domains = df["ipc_label"].str.strip().value_counts()
    covered      = set(your_domains.index.tolist())
    gaps         = [d for d in ALL_DOMAINS if d not in covered]

    col_w1, col_w2 = st.columns([2, 1])

    with col_w1:
        st.markdown("**Portfolio Coverage vs Technology Universe**")
        cover_data = []
        for domain in ALL_DOMAINS:
            count = your_domains.get(domain, 0)
            cover_data.append({
                "Domain":    domain,
                "Your Portfolio": count,
                "Status":    "✅ Covered" if count > 5 else ("⚠️ Thin" if count > 0 else "❌ Gap")
            })
        cover_df = pd.DataFrame(cover_data).sort_values("Your Portfolio", ascending=True)

        fig_cover = px.bar(
            cover_df, x="Your Portfolio", y="Domain",
            color="Status",
            color_discrete_map={"✅ Covered": "#1D9E75", "⚠️ Thin": "#BA7517", "❌ Gap": "#E24B4A"},
            orientation="h", height=520,
            labels={"Your Portfolio": "Patent Count", "Domain": ""}
        )
        fig_cover.update_layout(margin=dict(t=10, b=10),
                                yaxis=dict(autorange="reversed"),
                                legend=dict(orientation="h", y=-0.1))
        st.plotly_chart(fig_cover, use_container_width=True)

    with col_w2:
        st.markdown("**Whitespace Gaps**")
        st.error(f"**{len(gaps)} uncovered domains**")
        for g in gaps[:10]:
            st.markdown(f"❌ {g}")
        if len(gaps) > 10:
            st.caption(f"...and {len(gaps)-10} more")

        st.divider()
        st.markdown("**Thin Coverage (<5 patents)**")
        thin = [(d, c) for d, c in your_domains.items() if 0 < c < 5]
        for domain, cnt in thin[:8]:
            st.markdown(f"⚠️ {domain}: **{cnt}** patent{'s' if cnt>1 else ''}")

        st.divider()
        if st.button("🤖 Claude: Strategic Recommendations"):
            with st.spinner("Analysing gaps..."):
                gap_context = f"Covered domains: {list(covered)[:10]}\nGaps: {gaps[:10]}\nThin areas: {thin[:5]}"
                client = get_client()
                if client:
                    resp = client.messages.create(
                        model="claude-sonnet-4-6", max_tokens=600,
                        messages=[{"role": "user", "content":
                            f"As an IP strategist, give 5 specific patent filing recommendations "
                            f"based on these portfolio gaps. Be concrete and actionable.\n\n{gap_context}"}]
                    )
                    st.markdown(resp.content[0].text)

# ════════════════════════════════════════════════════════════════════════════════
# TAB 4 — PATENT Q&A CHAT
# ════════════════════════════════════════════════════════════════════════════════
with tab4:
    st.markdown("### Patent Portfolio Q&A")
    st.caption("Ask any question about your patent portfolio. Claude answers using semantic search over your 1,000 patents.")

    if "messages" not in st.session_state:
        st.session_state.messages = []

    # Example questions
    with st.expander("💡 Example questions"):
        examples = [
            "What are our strongest patents in video compression?",
            "Find patents related to DRAM memory with filing dates after 2018",
            "Which patents cover 5G beamforming technology?",
            "Summarize our semiconductor device portfolio",
            "What patents could be relevant to an M&A target in wireless communications?",
            "Which patents are most likely to expire in the next 3 years?",
            "Find prior art in our portfolio for OFDM signal processing",
        ]
        for ex in examples:
            if st.button(ex, key=f"ex_{ex[:20]}"):
                st.session_state.messages.append({"role": "user", "content": ex})

    # Display chat history
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    # Chat input
    if prompt := st.chat_input("Ask about your patent portfolio..."):
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)

        with st.chat_message("assistant"):
            with st.spinner("Searching patents + generating answer..."):
                results = search_patents(prompt, n=8)
                answer  = claude_answer(prompt, results)
                st.markdown(answer)

                with st.expander(f"📋 {len(results)} patents retrieved (sources)"):
                    for r in results:
                        score_color = "#1D9E75" if r['similarity'] > 0.7 else "#BA7517"
                        st.markdown(
                            f"**[{r.get('patent_id','')}]** {r.get('title','')}  \n"
                            f"IPC: {r.get('ipc_label','')} | Filed: {r.get('filing_date','')} | "
                            f"Status: {r.get('status','')} | "
                            f"<span style='color:{score_color}'>Similarity: {r['similarity']}</span>",
                            unsafe_allow_html=True
                        )

            st.session_state.messages.append({"role": "assistant", "content": answer})

    if st.button("🗑️ Clear chat"):
        st.session_state.messages = []
        st.rerun()

# ════════════════════════════════════════════════════════════════════════════════
# TAB 5 — REPORT GENERATOR
# ════════════════════════════════════════════════════════════════════════════════
with tab5:
    st.markdown("### AI Report Generator")
    st.caption("One-click Claude-written portfolio reports ready for executives, legal teams, or M&A discussions.")

    df = load_dataframe()

    col_r1, col_r2 = st.columns([1, 2])

    with col_r1:
        report_type = st.selectbox(
            "Report type",
            ["Portfolio Summary", "Landscape Report", "Whitespace Analysis",
             "M&A Diligence Summary", "Maintenance Recommendations"]
        )
        st.markdown("**Report covers:**")
        descriptions = {
            "Portfolio Summary":        "📊 KPIs · Tech focus · Filing velocity · Strategic recommendations",
            "Landscape Report":         "🗺️ Technology distribution · IPC analysis · Competitive positioning",
            "Whitespace Analysis":      "⬜ Coverage gaps · Filing opportunities · Risk areas",
            "M&A Diligence Summary":    "🤝 Portfolio quality · Geographic coverage · Value drivers · Risk",
            "Maintenance Recommendations": "💰 Cost optimization · Prune vs retain framework · Annuity strategy",
        }
        st.caption(descriptions.get(report_type, ""))

        custom_context = st.text_area(
            "Add context (optional)",
            placeholder="e.g., 'Focus on 5G patents', 'Target company is in wireless space'",
            height=80
        )

        generate_btn = st.button("⚡ Generate Report", type="primary", use_container_width=True)

    with col_r2:
        if generate_btn:
            with st.spinner(f"Claude writing {report_type}..."):
                report_text = claude_report(report_type, df)

            st.markdown(f"### {report_type}")
            st.markdown(report_text)
            st.divider()

            col_dl1, col_dl2 = st.columns(2)

            with col_dl1:
                st.download_button(
                    "📥 Download .txt",
                    report_text,
                    file_name=f"patent_{report_type.lower().replace(' ','_')}.txt",
                    mime="text/plain"
                )

            with col_dl2:
                try:
                    from docx import Document
                    from io import BytesIO
                    doc = Document()
                    doc.add_heading("PatentMind — " + report_type, 0)
                    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
                    doc.add_paragraph(f"Portfolio: {len(df)} patents")
                    doc.add_paragraph("")
                    for line in report_text.split("\n"):
                        if line.startswith("##"):
                            doc.add_heading(line.replace("#","").strip(), level=2)
                        elif line.startswith("#"):
                            doc.add_heading(line.replace("#","").strip(), level=1)
                        elif line.strip():
                            doc.add_paragraph(line)
                    buf = BytesIO()
                    doc.save(buf)
                    st.download_button(
                        "📥 Download .docx",
                        buf.getvalue(),
                        file_name=f"patent_{report_type.lower().replace(' ','_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except ImportError:
                    st.caption("Install python-docx for Word export: `pip install python-docx`")
        else:
            st.info("👈 Select a report type and click Generate Report")
            st.markdown("**Available report types:**")
            for rtype, desc in {
                "📊 Portfolio Summary": "High-level overview for executives — KPIs, tech focus, strategic direction",
                "🗺️ Landscape Report": "Technology distribution analysis — which domains you own and how strongly",
                "⬜ Whitespace Analysis": "Where to file next — gaps vs competitors, strategic opportunities",
                "🤝 M&A Diligence Summary": "Template for evaluating a target company's IP portfolio",
                "💰 Maintenance Recommendations": "Which patents to keep, prune, or monetize — cost optimization",
            }.items():
                st.markdown(f"**{rtype}**  \n{desc}\n")

# ── FOOTER ────────────────────────────────────────────────────────────────────
st.divider()
st.caption("PatentIntel POC · Built with Claude API · ChromaDB · Voyage-3 Embeddings · Streamlit · "
           "Free & open source · Powered by Anthropic")
